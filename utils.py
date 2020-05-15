import os
import ast
import docx
from bs4 import BeautifulSoup
import urllib.request as urllib2
import textract
import requests
import re

def LoadTxt(path):
    #words = []
    with open(path, encoding='utf8') as fr:
        words = fr.read()
    return words

def LoadDocx(path):
    document = docx.Document(path)
    doc = ""
    for para in document.paragraphs:
        doc += para.text + " "
    return doc

def LoadDoc(path):
    words = textract.process(path)
    words = words.decode('utf-8')
    return words

def LoadHtml(path):
    with open(path,'r') as f:
        html_doc = f.read()
    soup = BeautifulSoup(html_doc, 'html.parser')
    doc = ""
    for x in soup.find_all('p'):
        doc += x.text + " "
    return doc
def LoadUrl(url):
    result = requests.get(url)
    c = result.content
    soup = BeautifulSoup(c,"lxml")
    doc = ""
    for x in soup.find_all('p'):
        doc += x.text + " "
    return doc

def LoadDataFrom(path):
    regex = re.compile(r'((www\.[^\s]+)|(https?://[^\s]+))')
    match = regex.match(path)
    if match:
        return LoadUrl(path)
    else:
        _,file_extension = os.path.splitext(path)
        if file_extension == '.html':
            return LoadHtml(path)
        elif file_extension == '.docx':
            return LoadDocx(path)
        elif file_extension == '.doc':
            return LoadDoc(path)
        else:
            return LoadTxt(path)

def SaveTXT(path, text):
    f = open(path, 'w+')
    f.write(text)
    f.close()

def SaveDocx(path, text):
    document = docx.Document()
    document.add_paragraph(text)
    document.save(path)